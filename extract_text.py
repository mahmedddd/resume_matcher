import fitz  # PyMuPDF
import docx
import os

def extract_text_from_pdf(file):
    text = ""
    doc = fitz.open(stream=file.read(), filetype="pdf") if hasattr(file, "read") else fitz.open(file)
    for page in doc:
        text += page.get_text()
    return text.strip()

def extract_text_from_docx(file):
    if hasattr(file, "read"):  # Streamlit UploadedFile
        return "\n".join([para.text for para in docx.Document(file).paragraphs]).strip()
    else:
        return "\n".join([para.text for para in docx.Document(file).paragraphs]).strip()

def extract_text_from_txt(file):
    if hasattr(file, "read"):  # UploadedFile
        return file.read().decode("utf-8", errors="ignore").strip()
    else:  # Path
        with open(file, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()

def extract_resume_text(file):
    """Accepts either a file path or a Streamlit UploadedFile object."""
    ext = None

    if isinstance(file, str) or isinstance(file, os.PathLike):
        ext = os.path.splitext(file)[-1].lower()
    elif hasattr(file, "name"):  # UploadedFile
        ext = os.path.splitext(file.name)[-1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file)
    elif ext == ".docx":
        return extract_text_from_docx(file)
    elif ext == ".txt":
        return extract_text_from_txt(file)
    else:
        raise ValueError("Unsupported file type. Only PDF, DOCX, and TXT are supported.")
